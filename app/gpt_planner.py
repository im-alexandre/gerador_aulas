from __future__ import annotations

import json
import logging
import random
import time
from pathlib import Path
from typing import Any

from openai import OpenAI

from app.config.paths import APP_DIR, USER_INPUT_SLIDES
from app.debug_payload import dump_payload
from app.logging_utils import log_step
from app.prompt_utils import render_prompt_template

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

log = logging.getLogger(__name__)
MAX_RETRIES = 5
BASE_BACKOFF_SECONDS = 2


def _load_json_schema() -> dict[str, Any]:
    """
    Lê o schema de structured outputs para Responses API.

    Aceita:
      - wrapper: {"name": "...", "strict": true, "schema": {...}}
      - puro: {...}  (um JSON Schema)

    Retorna SEMPRE no formato:
      {"name": str, "strict": bool, "schema": {...}}
    """
    schema_path = APP_DIR / "prompts" / "schemas" / "slide_plan_v1.json"
    raw = json.loads(schema_path.read_text(encoding="utf-8"))

    if isinstance(raw, dict) and isinstance(raw.get("schema"), dict):
        name = str(raw.get("name") or "slide_plan_v1")
        strict = bool(raw.get("strict", True))
        schema = raw["schema"]
        return {"name": name, "strict": strict, "schema": schema}

    if isinstance(raw, dict) and raw.get("type"):
        return {"name": "slide_plan_v1", "strict": True, "schema": raw}

    raise ValueError(f"Schema inválido em {schema_path}")


def _safe_get(obj: Any, key: str, default: Any = None) -> Any:
    if obj is None:
        return default
    if isinstance(obj, dict):
        return obj.get(key, default)
    return getattr(obj, key, default)


def _try_parse_json_strict(s: str) -> dict[str, Any] | None:
    if not isinstance(s, str):
        return None
    s = s.strip()
    if not s:
        return None
    try:
        return json.loads(s)
    except Exception:
        return None


def _summarize_response_shapes(resp: Any) -> str:
    parts: list[str] = []
    out = _safe_get(resp, "output", []) or []
    parts.append(f"output_len={len(out)}")
    for i, msg in enumerate(out[:3]):
        parts.append(f"output[{i}].type={_safe_get(msg, 'type')}")
        content = _safe_get(msg, "content", []) or []
        parts.append(f" output[{i}].content_len={len(content)}")
        for j, c in enumerate(content[:5]):
            parts.append(f"  content[{j}].type={_safe_get(c, 'type')}")
    ot = _safe_get(resp, "output_text", None)
    parts.append(
        f"has_output_text={bool(ot)} output_text_len={len(ot) if isinstance(ot, str) else 0}"
    )
    return " | ".join(parts)


def docx_chars(path: Path) -> int:
    d = Document(path)
    paras = [p.text.strip() for p in d.paragraphs if p.text and p.text.strip()]
    return len("\n".join(paras))


def iter_block_items(parent):
    """Itera por parágrafos e tabelas mantendo a ordem no DOCX."""
    for child in parent.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def extract_docx_text(path: Path) -> str:
    """Extrai texto do DOCX preservando ordem de parágrafos e tabelas."""
    doc = Document(path)
    chunks: list[str] = []
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                chunks.append(text)
            continue
        for row in block.rows:
            cells = [cell.text.strip() for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                chunks.append(" | ".join(cells))
    return "\n".join(chunks).strip()


def with_backoff(fn, *args, **kwargs):
    """Executa uma função com backoff exponencial e jitter."""
    for attempt in range(1, MAX_RETRIES + 1):
        try:
            return fn(*args, **kwargs)
        except Exception as exc:
            if attempt == MAX_RETRIES:
                raise
            delay = min(BASE_BACKOFF_SECONDS * (2 ** (attempt - 1)), 30)
            delay += random.uniform(0, 0.5)
            log.warning(
                f"(with_backoff) ({fn}) Falha na API (tentativa {attempt}/{MAX_RETRIES}): {exc}"
            )
            log.warning(
                f"(with_backoff) Aguardando {delay:.1f}s antes de tentar novamente."
            )
            time.sleep(delay)


def upload_file(client: OpenAI, path: Path) -> str:
    """Faz upload de um arquivo para a API e retorna o file_id."""
    try:
        size = path.stat().st_size
    except OSError:
        size = -1

    log_step(
        log,
        path.parent.name,
        "upload_file",
        f"request: file={path.name} size_bytes={size} purpose=user_data",
        level=logging.DEBUG,
    )

    with open(path, "rb") as fh:
        f = with_backoff(client.files.create, file=fh, purpose="user_data")
    return f.id


def _extract_output_json(resp: Any, *, directory: str) -> dict[str, Any]:
    """
    Extrai o JSON estruturado da resposta do Responses.

    Preferência:
      1) content.type == "output_json" => retorna c.json
      2) content.type em ("output_text","text") => tenta json.loads no texto
      3) resp.output_text => tenta json.loads

    Se falhar, levanta erro com resumo do shape da resposta.
    """
    output = _safe_get(resp, "output", []) or []

    # 1) output_json direto
    for out in output:
        if _safe_get(out, "type") != "message":
            continue
        content = _safe_get(out, "content", []) or []
        for c in content:
            if _safe_get(c, "type") == "output_json":
                data = _safe_get(c, "json")
                if isinstance(data, dict):
                    return data
                try:
                    return dict(data)
                except Exception:
                    pass

    # 2) texto dentro do content (alguns SDKs devolvem output_text mesmo com schema)
    for out in output:
        if _safe_get(out, "type") != "message":
            continue
        content = _safe_get(out, "content", []) or []
        for c in content:
            ctype = _safe_get(c, "type")
            if ctype in ("output_text", "text"):
                # o campo pode ser "text" (string) ou {"text": "..."} dependendo do SDK
                txt = _safe_get(c, "text")
                if isinstance(txt, dict):
                    txt = txt.get("text") or txt.get("value")
                parsed = _try_parse_json_strict(txt if isinstance(txt, str) else "")
                if parsed is not None:
                    return parsed

    # 3) output_text do resp
    ot = _safe_get(resp, "output_text", None)
    parsed = _try_parse_json_strict(ot if isinstance(ot, str) else "")
    if parsed is not None:
        return parsed

    summary = _summarize_response_shapes(resp)
    log_step(
        log,
        directory,
        "_extract_output_json",
        f"response_shape: {summary}",
        level=logging.ERROR,
    )
    raise ValueError(
        "Resposta não retornou JSON parseável (nem output_json nem JSON em output_text)."
    )


def call_llm(
    client: OpenAI,
    model: str,
    instructions: str,
    file_ids: list[str],
    user_input: str,
    directory: str,
) -> dict[str, Any]:
    """Chama o modelo com arquivos anexados e retorna o JSON (dict)."""
    schema_fmt = _load_json_schema()

    tool_label = "code_interpreter"
    log_step(
        log,
        directory,
        "call_llm",
        (
            "request: "
            f"model={model} "
            f"files={len(file_ids)} "
            f"instructions_len={len(instructions or '')} "
            f"input_len={len(user_input or '')} "
            f"tool={tool_label} "
            f"json_schema={schema_fmt.get('name')}"
        ),
        level=logging.DEBUG,
    )

    payload: dict[str, Any] = {
        "model": model,
        "instructions": instructions,
        "input": user_input,
        "tools": [
            {
                "type": "code_interpreter",
                "container": {
                    "type": "auto",
                    "file_ids": file_ids,
                },
            }
        ],
        "tool_choice": {"type": "code_interpreter"},
        "text": {
            "format": {
                "type": "json_schema",
                "name": schema_fmt["name"],
                "strict": bool(schema_fmt.get("strict", True)),
                "schema": schema_fmt["schema"],
            }
        },
    }

    dump_path = dump_payload(payload)
    log_step(
        log, directory, "call_llm", f"request_dump={dump_path}", level=logging.DEBUG
    )

    resp = with_backoff(client.responses.create, **payload)
    return _extract_output_json(resp, directory=directory)


def generate_plan(
    client: OpenAI,
    prompt_md: str,
    content_docx: Path,
    roteiro_docx: Path,
    model: str,
    directory: str,
) -> dict[str, Any]:
    """Gera o plano de slides (JSON dict) a partir de conteúdo e roteiro."""
    content_id = upload_file(client, content_docx)
    roteiro_id = upload_file(client, roteiro_docx)
    file_ids = [content_id, roteiro_id]

    user_input = render_prompt_template(APP_DIR / USER_INPUT_SLIDES)

    log_step(
        log,
        directory,
        "upload_file",
        "Arquivos de conteudo + roteiro incluidos no pipeline",
    )
    log_step(log, directory, "call_llm", "LLM processando dados")

    return call_llm(
        client=client,
        model=model,
        instructions=prompt_md,
        file_ids=file_ids,
        directory=directory,
        user_input=user_input,
    )


def generate_plan_for_dir(
    api_key_override: str | None,
    content_docx: Path,
    roteiro_docx: Path,
    prompt_md: str,
    model: str,
    output_json: Path,
    force: bool = False,
    strict_json: bool = False,
    use_code_interpreter: bool = True,
) -> dict[str, Any] | None:
    """
    Gera e salva o JSON do plano para um diretório de núcleo.

    strict_json e use_code_interpreter são mantidos por compatibilidade com callers.
    """
    if output_json.exists() and not force:
        log_step(
            log,
            content_docx.parent.name,
            "generate_plan_for_dir",
            "Plano existente (reaproveitado)",
        )
        return None

    if api_key_override:
        api_key = api_key_override.strip()
    else:
        with open("app/prompts/openai_api_key") as key_file:
            api_key = key_file.read().strip()

    client = OpenAI(api_key=api_key)

    plan = generate_plan(
        client=client,
        prompt_md=prompt_md,
        content_docx=content_docx,
        roteiro_docx=roteiro_docx,
        model=model,
        directory=content_docx.parent.name,
    )

    log_step(
        log,
        content_docx.parent.name,
        "generate_plan_for_dir",
        "Formulando abstracao dos slides",
    )

    output_json.write_text(
        json.dumps(plan, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    log_step(
        log,
        content_docx.parent.name,
        "generate_plan_for_dir",
        f"Plano gerado: {output_json}",
        level=logging.DEBUG,
    )

    return plan
