from __future__ import annotations

import logging
import re
import shutil
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph


log = logging.getLogger(__name__)

HEADING_MODULE = "Heading 1"
HEADING_NUCLEUS = "Heading 2"


def iter_block_items(doc: DocxDocument) -> Iterable[Paragraph | Table]:
    """Itera sobre parágrafos e tabelas na ordem em que aparecem no documento."""
    for child in doc.element.body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, doc)
        elif child.tag.endswith("}tbl"):
            yield Table(child, doc)


def detect_nucleus_kind(title: str) -> str:
    """Infere se o núcleo é conceitual (c) ou prático (p) pelo título."""
    title_upper = title.upper()
    if "PRÁTICO" in title_upper or "PRATICO" in title_upper or "NP" in title_upper:
        return "p"
    if "CONCEITUAL" in title_upper or "NC" in title_upper:
        return "c"
    return "c"


def extract_first_int(text: str) -> int | None:
    """Extrai o primeiro número inteiro presente no texto."""
    match = re.search(r"\d+", text)
    return int(match.group(0)) if match else None


def clear_headers_footers(doc: DocxDocument) -> None:
    """Remove conteúdo de cabeçalho e rodapé."""
    for section in doc.sections:
        for part in (section.header, section.footer):
            for paragraph in list(part.paragraphs):
                paragraph._element.getparent().remove(paragraph._element)
            for table in list(part.tables):
                table._element.getparent().remove(table._element)


def split_docx_to_nuclei(docx_path: Path, output_root: Path, force: bool) -> list[Path]:
    """Divide um DOCX em núcleos conforme Heading 1/2."""
    doc = Document(str(docx_path))
    blocks = list(iter_block_items(doc))
    module_idx = 0
    nucleus_counts = {"c": 0, "p": 0}
    output_docs: list[Path] = []

    current_start: int | None = None
    current_meta: tuple[int, str, int] | None = None
    segments: list[tuple[int, int, int, str, int]] = []

    def finalize(end_idx: int) -> None:
        nonlocal current_start, current_meta
        if current_start is None or current_meta is None:
            return
        mod, kind, number = current_meta
        segments.append((current_start, end_idx, mod, kind, number))
        current_start = None
        current_meta = None

    for idx, block in enumerate(blocks):
        if not isinstance(block, Paragraph):
            continue
        style_name = block.style.name if block.style else ""
        text = (block.text or "").strip()

        if style_name == HEADING_MODULE:
            finalize(idx)
            mod_number = extract_first_int(text)
            module_idx = mod_number if mod_number is not None else module_idx + 1
            nucleus_counts = {"c": 0, "p": 0}
            continue

        if style_name == HEADING_NUCLEUS:
            finalize(idx)
            if module_idx == 0:
                module_idx = 1

            kind = detect_nucleus_kind(text)
            number = extract_first_int(text)
            if number is None:
                nucleus_counts[kind] += 1
                number = nucleus_counts[kind]
            else:
                nucleus_counts[kind] = max(nucleus_counts[kind], number)

            current_start = idx
            current_meta = (module_idx, kind, number)

    finalize(len(blocks))

    if not segments:
        return output_docs

    for start, end, mod, kind, number in segments:
        nucleus_name = f"mod{mod}_n{kind}{number}"
        nucleus_dir = output_root / nucleus_name
        nucleus_dir.mkdir(parents=True, exist_ok=True)
        docx_out = nucleus_dir / f"{nucleus_name}.docx"

        if docx_out.exists() and not force:
            log.info(f"[splitter] Mantendo existente: {docx_out.name}")
            continue

        shutil.copyfile(docx_path, docx_out)
        out_doc = Document(str(docx_out))
        clear_headers_footers(out_doc)

        out_blocks = list(iter_block_items(out_doc))
        for idx in range(len(out_blocks) - 1, -1, -1):
            if idx < start or idx >= end:
                blk = out_blocks[idx]
                blk._element.getparent().remove(blk._element)

        out_doc.save(str(docx_out))
        output_docs.append(docx_out)

    return output_docs


def split_course_content(course_dir: Path, force: bool) -> list[Path]:
    """Encontra DOCX na raiz do curso e divide em núcleos."""
    docxs = [
        p
        for p in course_dir.glob("*.docx")
        if not p.name.startswith("ROT_") and not p.name.endswith("_tagged.docx")
    ]
    if not docxs:
        return []

    docxs = sorted(docxs)
    created: list[Path] = []

    # VIDINT usa o material inteiro do curso.
    created += create_vidint_docx(docxs[0], course_dir, force=force)

    if len(docxs) > 1:
        log.warning(
            "[splitter] Mais de um DOCX encontrado na raiz; usando o primeiro para VIDINT."
        )

    for docx_path in docxs:
        created += split_docx_to_nuclei(docx_path, course_dir, force=force)
    return created


def create_vidint_docx(
    docx_path: Path,
    output_root: Path,
    force: bool,
) -> list[Path]:
    """Cria mod0_vidint com o conteúdo completo do curso."""
    nucleus_name = "mod0_vidint"
    nucleus_dir = output_root / nucleus_name
    nucleus_dir.mkdir(parents=True, exist_ok=True)
    docx_out = nucleus_dir / f"{nucleus_name}.docx"

    if docx_out.exists() and not force:
        log.info(f"[splitter] Mantendo existente: {docx_out.name}")
        return []

    shutil.copyfile(docx_path, docx_out)
    out_doc = Document(str(docx_out))
    clear_headers_footers(out_doc)
    out_doc.save(str(docx_out))
    return [docx_out]
