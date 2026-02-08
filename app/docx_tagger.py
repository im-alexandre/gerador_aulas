from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


NS = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
}


EXT_BY_CONTENT_TYPE = {
    "image/png": ".png",
    "image/jpeg": ".jpg",
    "image/jpg": ".jpg",
    "image/gif": ".gif",
    "image/bmp": ".bmp",
    "image/tiff": ".tif",
    "image/webp": ".webp",
}


def iter_paragraphs(doc: Document) -> Iterable[Paragraph]:
    """Itera por todos os parágrafos do documento, incluindo tabelas aninhadas."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        yield from iter_paragraphs_in_table(table)


def iter_paragraphs_in_table(table: Table) -> Iterable[Paragraph]:
    """Itera por parágrafos dentro de uma tabela, incluindo tabelas aninhadas."""
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p
            for t in cell.tables:
                yield from iter_paragraphs_in_table(t)


def replace_run_with_text(run, text: str) -> None:
    """Substitui o conteúdo de um run por um texto simples."""
    r = run._element
    for child in list(r):
        r.remove(child)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)


def guess_extension(image_part) -> str:
    """Inferre a extensão do arquivo de imagem pelo nome ou content-type."""
    filename = getattr(image_part, "filename", "") or ""
    if filename:
        return Path(filename).suffix or ".png"
    content_type = getattr(image_part, "content_type", "") or ""
    return EXT_BY_CONTENT_TYPE.get(content_type, ".png")


def tag_images_in_docx(
    docx_path: Path,
    assets_dir: Path,
    tag_prefix: str = "assets",
) -> int:
    """Extrai imagens do DOCX, grava em assets e substitui imagens por tags."""
    assets_dir.mkdir(parents=True, exist_ok=True)
    doc = Document(str(docx_path))
    tag_prefix = tag_prefix.strip("/\\")

    img_index = 1
    for paragraph in iter_paragraphs(doc):
        for run in paragraph.runs:
            blips = run._element.xpath(".//*[local-name()='blip']")
            if not blips:
                continue
            tags: list[str] = []
            for blip in blips:
                rel_id = blip.get(qn("r:embed"))
                if not rel_id:
                    continue
                image_part = doc.part.related_parts.get(rel_id)
                if not image_part:
                    continue
                ext = guess_extension(image_part)
                filename = f"img_{img_index:04d}{ext}"
                img_index += 1
                (assets_dir / filename).write_bytes(image_part.blob)
                tag_path = f"{tag_prefix}/{filename}" if tag_prefix else filename
                tags.append(f"[[IMG:{tag_path}]]")
            if tags:
                replace_run_with_text(run, " ".join(tags))

    doc.save(str(docx_path))
    return img_index - 1


def find_content_docx(directory: Path) -> Path | None:
    """Encontra o DOCX de conteúdo em um diretório de núcleo."""
    named = directory / f"{directory.name}.docx"
    if named.exists():
        return named
    candidates = [
        p
        for p in directory.glob("*.docx")
        if not re.match(r"^ROT_", p.name) and not p.name.endswith("_tagged.docx")
    ]
    return candidates[0] if candidates else None


def create_tagged_docx(
    source_docx: Path,
    tagged_docx: Path,
    assets_dir: Path,
    tag_prefix: str,
) -> int:
    """Gera o DOCX tagueado a partir do original."""
    tagged_docx.parent.mkdir(parents=True, exist_ok=True)
    tagged_docx.write_bytes(source_docx.read_bytes())
    return tag_images_in_docx(
        docx_path=tagged_docx,
        assets_dir=assets_dir,
        tag_prefix=tag_prefix,
    )


def find_roteiro_docx(directory: Path) -> Path | None:
    """Encontra o DOCX de roteiro (ROT_*) em um diretório de núcleo."""
    candidates = list(directory.glob("ROT_*.docx"))
    return candidates[0] if candidates else None
